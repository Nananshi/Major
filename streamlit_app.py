import os
import sys
import tempfile
import importlib.util
from collections import Counter
from docx import Document
import io
import streamlit as st
import pandas as pd

# Add 2.o folder to path for imports
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '2.o'))


# NOTE: heavy modules (spaCy, PyMuPDF) are imported lazily to avoid startup delays


def load_module_from_path(path, name="module"):
    spec = importlib.util.spec_from_file_location(name, path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


st.set_page_config(page_title="Resume Enhancer", layout="wide")

st.title("Resume & JD Enhancer")
st.markdown("### Advanced Resume Matching & Enhancement using Local ML Pipeline")

# Sidebar: explanation of the pipeline
with st.sidebar:
    st.header("ℹ️ Pipeline Info")
    st.markdown("""
    **Local Processing Pipeline:**
    - Text Extraction (PDF/TXT)
    - Normalization & Tokenization
    - Stop Word Removal
    - Named Entity Recognition (NER)
    - Parsing (spaCy)
    - Resume-JD Matching (Cosine Similarity)
    - Skill Gap Analysis

    No external API calls required! Uses your local ML models.
    """)

st.header("Inputs")
col1, col2 = st.columns(2)

with col1:
    st.subheader("Resume (File or Text)")
    resume_file = st.file_uploader("Upload resume (PDF or TXT)", type=["pdf", "txt"], key="resume")
    resume_text_area = st.text_area("Or paste resume text here")

with col2:
    st.subheader("Job Description (File or Text)")
    jd_file = st.file_uploader("Upload JD (PDF or TXT)", type=["pdf", "txt"], key="jd")
    jd_text_area = st.text_area("Or paste JD text here")


def extract_text_from_uploaded(uploaded):
    if uploaded is None:
        return ""
    name = uploaded.name.lower()
    data = uploaded.getvalue()
    if name.endswith('.pdf'):
        try:
            import pdfplumber
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(data)
                tmp.flush()
                with pdfplumber.open(tmp.name) as pdf:
                    text = "\n".join([p.extract_text() or "" for p in pdf.pages])
            return text
        except Exception as e:
            st.warning(f"PDF extraction failed: {e}")
            return ""
    else:
        try:
            return data.decode('utf-8')
        except Exception:
            return str(data)


resume_text = resume_text_area.strip() if resume_text_area.strip() else extract_text_from_uploaded(resume_file)
jd_text = jd_text_area.strip() if jd_text_area.strip() else extract_text_from_uploaded(jd_file)

if not resume_text:
    st.warning("Please provide a resume text or upload a file to continue.")

run_btn = st.button("Run Enhancement")

if run_btn and resume_text:
    st.info("Running local preprocessing pipeline — this may take a moment (spaCy models loading)...")

    # ========== LAZY IMPORTS (heavy modules) ==========
    try:
        from normalization import normalize_text as norm_text_func
        from tokenization import tokenize_text as tokenize_func
        from stop_word_removal import remove_stopwords as remove_stopwords_func
        from NER import perform_full_ner
        from parsing import parse_tokens
        import tempfile as tf
        import ast
    except Exception as e:
        st.error(f"Failed to import 2.o modules: {e}")
        st.stop()

    # ========== PIPELINE EXECUTION ==========

    # Step 1: Extract or use text as-is
    resume_text_extracted = resume_text
    jd_text_extracted = jd_text if jd_text else ""

    if not jd_text_extracted:
        st.warning("⚠️ Job Description not provided. Proceeding with Resume analysis only.")

    # Step 2: Normalize both texts
    try:
        resume_normalized = norm_text_func(resume_text_extracted)
        jd_normalized = norm_text_func(jd_text_extracted) if jd_text_extracted else ""
        st.success("✅ Normalization complete")
    except Exception as e:
        st.error(f"Normalization failed: {e}")
        st.stop()

    # Step 3: Tokenize
    try:
        resume_tokens = tokenize_func(resume_normalized)
        jd_tokens = tokenize_func(jd_normalized) if jd_normalized else []
        st.success("✅ Tokenization complete")
    except Exception as e:
        st.error(f"Tokenization failed: {e}")
        st.stop()

    # Step 4: Remove stopwords
    try:
        resume_tokens_str = " ".join(resume_tokens)
        resume_no_stop = remove_stopwords_func(resume_tokens_str).split()

        jd_tokens_str = " ".join(jd_tokens) if jd_tokens else ""
        jd_no_stop = remove_stopwords_func(jd_tokens_str).split() if jd_tokens_str else []

        st.success("✅ Stop word removal complete")
    except Exception as e:
        st.error(f"Stop word removal failed: {e}")
        st.stop()

    # Step 5: Perform NER on resume tokens (write to temp file first)
    try:
        resume_ner_output = None
        jd_ner_output = None

        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f_resume:
            f_resume.write(str(resume_tokens))
            resume_tokens_file = f_resume.name

        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f_jd:
            f_jd.write(str(jd_tokens))
            jd_tokens_file = f_jd.name

        # Perform NER (returns structured dict)
        resume_ner_output = perform_full_ner(resume_tokens_file, output_dir=tempfile.gettempdir())
        if jd_tokens:
            jd_ner_output = perform_full_ner(jd_tokens_file, output_dir=tempfile.gettempdir())

        st.success("✅ NER complete")
    except Exception as e:
        st.warning(f"⚠️ NER processing had issues: {e}")
        resume_ner_output = {}
        jd_ner_output = {}

    # ========== VISUALIZATIONS ==========

    st.header("📊 Analysis & Insights")

    # Tab 1: Preprocessing Overview
    with st.expander("Preprocessing Overview", expanded=True):
        col1, col2, col3 = st.columns(3)

        with col1:
            st.metric("Resume Tokens", len(resume_tokens))
            st.metric("After Stopword Removal", len(resume_no_stop))

        with col2:
            if jd_tokens:
                st.metric("JD Tokens", len(jd_tokens))
                st.metric("JD After Stopword Removal", len(jd_no_stop))

        with col3:
            st.metric("Unique Resume Terms", len(set(resume_no_stop)))
            if jd_no_stop:
                st.metric("Unique JD Terms", len(set(jd_no_stop)))

    # Tab 2: Token Frequency Analysis
    with st.expander("Token Frequency Analysis"):
        col1, col2 = st.columns(2)

        with col1:
            st.subheader("Resume - Top Tokens")
            counts = Counter([t.lower() for t in resume_no_stop])
            df_counts = pd.DataFrame(counts.most_common(20), columns=["token", "count"])
            if not df_counts.empty:
                st.bar_chart(df_counts.set_index("token"))
            else:
                st.info("No tokens found.")

        with col2:
            if jd_no_stop:
                st.subheader("JD - Top Tokens")
                counts_jd = Counter([t.lower() for t in jd_no_stop])
                df_counts_jd = pd.DataFrame(counts_jd.most_common(20), columns=["token", "count"])
                if not df_counts_jd.empty:
                    st.bar_chart(df_counts_jd.set_index("token"))
                else:
                    st.info("No tokens found.")

    # Tab 3: Named Entities
    with st.expander("Named Entities & Keywords"):
        col1, col2 = st.columns(2)

        with col1:
            st.subheader("Resume - Extracted Entities")
            if resume_ner_output:
                for entity_type, entities in resume_ner_output.items():
                    if entities:
                        st.markdown(f"**{entity_type}**: {', '.join(entities[:10])}")
                        if len(entities) > 10:
                            st.caption(f"... and {len(entities) - 10} more")
            else:
                st.info("No entities extracted.")

        with col2:
            st.subheader("JD - Extracted Entities")
            if jd_ner_output:
                for entity_type, entities in jd_ner_output.items():
                    if entities:
                        st.markdown(f"**{entity_type}**: {', '.join(entities[:10])}")
                        if len(entities) > 10:
                            st.caption(f"... and {len(entities) - 10} more")
            else:
                st.info("No entities extracted or JD not provided.")

    # Tab 4: Resume-JD Matching (using local similarity module)
    if jd_text_extracted:
        with st.expander("Resume-JD Match Score", expanded=True):
            try:
                # Import the local matching module
                from Similarity.Resume_JD_Matching import match_resume_jd, load_ner_data

                # Save NER data to temp files for matching
                with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
                    for k, v in resume_ner_output.items():
                        f.write(f"{k}: {v}\n")
                    resume_ner_file = f.name

                with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
                    for k, v in jd_ner_output.items():
                        f.write(f"{k}: {v}\n")
                    jd_ner_file = f.name

                # Get match scores and overall match
                match_scores, overall_match = match_resume_jd(resume_ner_file, jd_ner_file)

                # Display Overall Match prominently at top
                st.markdown("### Overall Match Score")

                # Determine color based on score
                if overall_match >= 70:
                    color = "green"
                    status = "Excellent Match! ✅"
                elif overall_match >= 50:
                    color = "orange"
                    status = "Good Match 👍"
                else:
                    color = "red"
                    status = "Needs Improvement ⚠️"

                # Display overall match with large metric
                col_center = st.columns([1, 2, 1])[1]
                with col_center:
                    st.markdown(f"""
                    <div style="
                        text-align: center; 
                        padding: 20px; 
                        background-color: {'#d4edda' if color == 'green' else '#fff3cd' if color == 'orange' else '#f8d7da'}; 
                        border-radius: 10px;
                        border: 2px solid {color};
                    ">
                        <h1 style="color: {color}; margin: 0;">{overall_match:.1f}%</h1>
                        <p style="margin: 5px 0 0 0; font-size: 18px; color: {color};">{status}</p>
                    </div>
                    """, unsafe_allow_html=True)

                st.markdown("---")

                # Display individual field scores (removed Achievements)
                st.markdown("### 📊 Field-wise Match Scores")

                col1, col2, col3, col4 = st.columns(4)

                with col1:
                    skills_score = match_scores.get('SKILLS', 0)
                    st.metric("Skills", f"{skills_score:.1f}%")

                with col2:
                    exp_score = match_scores.get('EXPERIENCE', 0)
                    st.metric("Experience", f"{exp_score:.1f}%")

                with col3:
                    edu_score = match_scores.get('EDUCATION', 0)
                    st.metric("Education", f"{edu_score:.1f}%")

                with col4:
                    proj_score = match_scores.get('PROJECTS', 0)
                    st.metric("Projects", f"{proj_score:.1f}%")

                st.markdown("---")

                # Skill Gaps Analysis
                st.markdown("### Skill Gap Analysis")
                cv_skills = set(resume_ner_output.get("SKILLS", []))
                jd_skills = set(jd_ner_output.get("SKILLS", []))
                missing_skills = jd_skills - cv_skills
                matched_skills = cv_skills & jd_skills

                col1, col2 = st.columns(2)

                with col1:
                    st.markdown("#### ✅ Matched Skills")
                    if matched_skills:
                        st.success(f"Found {len(matched_skills)} matching skills:")
                        for skill in sorted(matched_skills):
                            st.write(f"  • {skill}")
                    else:
                        st.info("No exact skill matches found.")

                with col2:
                    st.markdown("#### ⚠️ Missing Skills")
                    if missing_skills:
                        st.warning(f"Skills in JD but missing from Resume ({len(missing_skills)}):")
                        for skill in sorted(missing_skills):
                            st.write(f"  • {skill}")
                    else:
                        st.success("✅ All JD skills are present in Resume!")

                # Detailed breakdown table (removed Achievements)
                st.markdown("---")
                st.markdown("### 📋 Detailed Score Breakdown")

                score_data = []
                # Only include these 4 fields
                for field in ['SKILLS', 'EXPERIENCE', 'EDUCATION', 'PROJECTS']:
                    score = match_scores.get(field, 0)
                    if score >= 70:
                        rating = "Very Good"
                        emoji = "✅"
                    elif score >= 50:
                        rating = "Good"
                        emoji = "👍"
                    elif score >= 30:
                        rating = "Ok"
                        emoji = "⚠️"
                    else:
                        rating = "Poor"
                        emoji = "❌"

                    score_data.append({
                        "Field": field,
                        "Score": f"{score:.2f}%",
                        "Rating": rating,
                        "Status": emoji
                    })

                df_scores = pd.DataFrame(score_data)
                st.dataframe(df_scores, use_container_width=True, hide_index=True)

            except Exception as e:
                st.error(f"❌ Resume-JD matching failed: {e}")
                st.exception(e)  # Show full error for debugging

    # Tab 5: Enhanced Resume Suggestions
    with st.expander("Resume Enhancement Suggestions"):
        st.subheader("Recommended Changes")

        suggestions = []

        # Suggestion 1: Missing skills
        if jd_text_extracted:
            cv_skills = set(resume_ner_output.get("SKILLS", []))
            jd_skills = set(jd_ner_output.get("SKILLS", []))
            missing_skills = jd_skills - cv_skills
            if missing_skills:
                suggestions.append(f"Add these skills to your resume: {', '.join(sorted(list(missing_skills)[:5]))}")

        # Suggestion 2: Token diversity
        if len(resume_no_stop) < 50:
            suggestions.append(" Consider expanding your resume with more details and keywords.")
        elif len(resume_no_stop) > 500:
            suggestions.append(" Your resume might be too long. Consider making it more concise.")

        # Suggestion 3: Keyword optimization
        common_keywords = set(resume_no_stop) & set(jd_no_stop) if jd_no_stop else set()
        if common_keywords:
            suggestions.append(f" You have {len(common_keywords)} keywords matching the JD. Good alignment!")

        if suggestions:
            for i, sug in enumerate(suggestions, 1):
                st.markdown(f"{i}. {sug}")
        else:
            st.info("No specific suggestions at this time.")

    # Tab 6: Download Enhanced Resume
    with st.expander("Generate Enhanced Resume"):
        st.subheader("📄 Download Your Enhanced Resume")

        # Create a simple enhanced version by combining skills
        enhanced_text = f"""ENHANCED RESUME

{resume_text}

---

ADDITIONAL RECOMMENDED SKILLS (from JD matching):
"""

        if jd_text_extracted and jd_ner_output:
            cv_skills = set(resume_ner_output.get("SKILLS", []))
            jd_skills = set(jd_ner_output.get("SKILLS", []))
            missing_skills = sorted(list(jd_skills - cv_skills))
            if missing_skills:
                enhanced_text += "\n".join(missing_skills[:15])
            else:
                enhanced_text += "None - your resume skills align perfectly with the JD!"

        # Display enhanced text in a text area
        st.text_area("Enhanced Resume Preview", enhanced_text, height=300)


        # Function to generate PDF
        def generate_word_resume(text):
            try:
                doc = Document()

                # Add content line by line
                for line in text.split("\n"):
                    line = line.strip()
                    if not line:
                        doc.add_paragraph("")  # blank line
                        continue

                    # Section headers (uppercase or with ---)
                    if line.isupper() or line.startswith("---"):
                        p = doc.add_paragraph()
                        run = p.add_run(line)
                        run.bold = True
                    else:
                        doc.add_paragraph(line)

                # Save file to a bytes buffer
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                return buffer.getvalue()

            except Exception as e:
                st.error(f"Word file generation failed: {e}")
                return None


        # Create two columns for download buttons
        col1, col2 = st.columns(2)

        with col1:
            st.download_button(
                label="📥 Download as TXT",
                data=enhanced_text.encode("utf-8"),
                file_name="enhanced_resume.txt",
                mime="text/plain"
            )

        with col2:
            word_data = generate_word_resume(enhanced_text)
            if word_data:
                st.download_button(
                    label="📄 Download as Word (.docx)",
                    data=word_data,
                    file_name="enhanced_resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )